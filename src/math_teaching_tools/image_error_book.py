from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.shared import Inches


IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg"}


def iter_images(image_folder: Path) -> list[Path]:
    return sorted(
        path
        for path in image_folder.iterdir()
        if path.is_file() and path.suffix.lower() in IMAGE_EXTENSIONS
    )


def create_error_book(image_folder: Path, output_docx: Path, image_width: float) -> None:
    document = Document()
    document.add_heading("Math Error Book", level=0)

    images = iter_images(image_folder)
    if not images:
        document.add_paragraph("No problem images found.")

    for index, image_path in enumerate(images, start=1):
        document.add_heading(f"Problem {index}: {image_path.name}", level=1)
        document.add_picture(str(image_path), width=Inches(image_width))
        document.add_paragraph("Notes:")
        document.add_paragraph("")

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_docx)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Create a Word error book from a folder of problem images.",
    )
    parser.add_argument("image_folder", type=Path)
    parser.add_argument("output_docx", type=Path)
    parser.add_argument("--image-width", type=float, default=4.0)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    create_error_book(args.image_folder, args.output_docx, args.image_width)
    print(f"Wrote error book: {args.output_docx}")


if __name__ == "__main__":
    main()
